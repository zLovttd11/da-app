"""Generate a comprehensive Word (.docx) analysis report with executive summary, methodology, findings, and recommendations."""

import os
import io
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import pandas as pd

OUTPUT_DIR = Path(__file__).parent.parent / "outputs"


def _add_image_from_bytes(doc, png_bytes, width_inches=5.0):
    if not png_bytes:
        return
    try:
        stream = io.BytesIO(png_bytes)
        doc.add_picture(stream, width=Inches(width_inches))
        doc.add_paragraph()
    except Exception:
        pass


def _add_dataframe_table(doc, pdf, col_widths=None):
    if pdf.empty:
        doc.add_paragraph("No data available.")
        return
    table = doc.add_table(rows=len(pdf) + 1, cols=len(pdf.columns), style="Light Grid Accent 1")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, col_name in enumerate(pdf.columns):
        cell = table.rows[0].cells[j]
        cell.text = str(col_name)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(8)
    for i, (_, row) in enumerate(pdf.iterrows()):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            if isinstance(val, float):
                cell.text = "{:.4f}".format(val) if abs(val) < 100 else "{:.2f}".format(val)
            else:
                cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
    doc.add_paragraph()


def _make_para(doc, text, bold=False, size=10, color=None, alignment=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if alignment:
        p.alignment = alignment
    return p


def _make_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h


def generate_report(
    dataset_name, df, col_types, profile, desc_stats, cat_summary,
    static_charts, corr_result, reg_result, hyp_result,
    class_result=None, cluster_result=None, fi_result=None,
    reg_compare_result=None, split_result=None, target_col=None,
    analysis_mode="auto", auto_ml_result=None, preprocess_log=None,
    problem_statement="",
):
    """Generate the full Word report."""
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    # ===== COVER PAGE =====
    for _ in range(6):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Data Analysis Report")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Dataset: {}".format(dataset_name))
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    if problem_statement:
        problem_para = doc.add_paragraph()
        problem_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        short_problem = problem_statement[:200] + ("..." if len(problem_statement) > 200 else "")
        run = problem_para.add_run(short_problem)
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run("Generated: {}".format(datetime.now().strftime("%B %d, %Y")))
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    tool_para = doc.add_paragraph()
    tool_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tool_para.add_run("Autonomous Analysis by DA App v2")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    doc.add_page_break()

    # ===== 1. EXECUTIVE SUMMARY =====
    _make_heading(doc, "1. Executive Summary", level=1)
    rows, cols = profile.get("shape", (0, 0))
    n_numeric = len(col_types.get("numeric", []))
    n_cat = len(col_types.get("categorical", []))
    missing = profile.get("total_missing_cells", 0)
    dupes = profile.get("duplicate_rows", 0)

    summary_parts = [
        'This report presents a comprehensive data analysis of the "{}" dataset, containing {} observations across {} variables ({} numeric, {} categorical).'.format(
            dataset_name, rows, cols, n_numeric, n_cat),
    ]

    if reg_result and reg_result.get("success"):
        summary_parts.append(
            "Regression analysis identified key predictors of {} with an R-squared of {:.3f}.".format(
                target_col, reg_result["r_squared"]))

    if class_result and class_result.get("success"):
        summary_parts.append(
            "Classification analysis compared {} models; the best performer ({}) achieved an F1 score of {:.3f} across {} classes.".format(
                len(class_result.get("metrics", [])), class_result.get("best_model", "N/A"),
                class_result["metrics"][0].get("f1_score", 0) if class_result["metrics"] else 0,
                class_result.get("n_classes", 2)))

    if hyp_result:
        summary_parts.append(
            "{} out of {} hypothesis tests yielded statistically significant results.".format(
                hyp_result.get("significant_count", 0), hyp_result.get("total_tests", 0)))

    if cluster_result and cluster_result.get("success"):
        summary_parts.append(
            "Clustering analysis identified {} distinct segments (silhouette score: {:.3f}).".format(
                cluster_result["best_k"], max(s["silhouette"] for s in cluster_result["silhouette_scores"])))

    summary_parts.append(
        "The dataset contains {} missing values and {} duplicate rows. Key findings and recommendations are detailed in the sections below.".format(
            missing, dupes))

    for part in summary_parts:
        _make_para(doc, part, size=10)

    doc.add_page_break()

    # ===== 2. DATA OVERVIEW =====
    _make_heading(doc, "2. Data Overview", level=1)
    _make_heading(doc, "2.1 Dataset Summary", level=2)

    summary_data = [("Observations", str(rows)), ("Variables", str(cols)),
                    ("Missing Cells", str(missing)), ("Duplicate Rows", str(dupes)),
                    ("Memory Usage", str(profile.get("memory_usage", "N/A"))),
                    ("Target Variable", target_col if target_col else "None")]
    table = doc.add_table(rows=len(summary_data), cols=2, style="Light Grid Accent 1")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(summary_data):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
        for cell in table.rows[i].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()

    _make_heading(doc, "2.2 Column Types", level=2)
    type_data = [("Numeric", str(n_numeric)), ("Categorical", str(n_cat)),
                 ("Datetime", str(len(col_types.get("datetime", [])))),
                 ("Text/Other", str(len(col_types.get("text", []))))]
    table = doc.add_table(rows=len(type_data) + 1, cols=2, style="Light Grid Accent 1")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.rows[0].cells[0].text = "Type"
    table.rows[0].cells[1].text = "Count"
    for i, (k, v) in enumerate(type_data):
        table.rows[i + 1].cells[0].text = k
        table.rows[i + 1].cells[1].text = v
    doc.add_paragraph()

    _make_heading(doc, "2.3 Data Quality", level=2)
    missing_info = profile.get("missing", {})
    if missing_info:
        miss_df = pd.DataFrame([{"Variable": c, "Missing Count": info["count"],
                                  "Missing %": "{}%".format(info["percent"])}
                                 for c, info in missing_info.items() if info["count"] > 0])
        if not miss_df.empty:
            _add_dataframe_table(doc, miss_df)
        else:
            _make_para(doc, "No missing values detected.", size=10)
    doc.add_page_break()

    # ===== 3. METHODOLOGY =====
    _make_heading(doc, "3. Methodology", level=1)
    _make_para(doc, "The analysis follows the CRISP-DM framework and proceeds through multiple phases, "
               "building from descriptive exploration to predictive modeling and clustering.", size=10)

    phases = [
        ("Phase 1: Descriptive Exploration", "Summary statistics, distribution analysis, and data profiling."),
        ("Phase 2: Correlation Analysis", "Pearson correlation matrix and pairwise relationships."),
        ("Phase 3: Regression Modeling",
         "OLS linear regression with residual diagnostics and multi-model comparison."),
        ("Phase 4: Classification Analysis",
         "Multi-model comparison (Logistic Regression, Random Forest, Gradient Boosting) with cross-validation."),
        ("Phase 5: Hypothesis Testing",
         "Independent t-tests, one-way ANOVA, and chi-square tests with effect sizes."),
        ("Phase 6: Feature Engineering",
         "Random Forest importance and mutual information to identify key predictors."),
        ("Phase 7: Clustering",
         "K-means clustering with silhouette analysis and PCA visualization."),
    ]
    for title, desc in phases:
        _make_para(doc, title, bold=True, size=10)
        _make_para(doc, desc, size=9)

    doc.add_page_break()

    # ===== 4. DESCRIPTIVE ANALYSIS =====
    _make_heading(doc, "4. Descriptive Analysis", level=1)
    if not desc_stats.empty:
        _make_heading(doc, "4.1 Numeric Variables", level=2)
        _add_dataframe_table(doc, desc_stats)
    if cat_summary:
        _make_heading(doc, "4.2 Categorical Variables", level=2)
        for col, freq_df in cat_summary.items():
            _make_para(doc, "Distribution of '{}':".format(col), bold=True, size=10)
            _add_dataframe_table(doc, freq_df.head(15))
    doc.add_page_break()

    # ===== 5. VISUALIZATIONS =====
    _make_heading(doc, "5. Data Visualizations", level=1)
    chart_sections = [
        ("5.1 Distribution Histograms", "histograms"),
        ("5.2 Box Plot Comparison", "boxplot"),
        ("5.3 Categorical Bar Charts", "barcharts"),
    ]
    for title, key in chart_sections:
        if static_charts.get(key):
            _make_heading(doc, title, level=2)
            for png in static_charts[key]:
                _add_image_from_bytes(doc, png, 5.0)
    doc.add_page_break()

    # ===== 6. CORRELATION =====
    _make_heading(doc, "6. Correlation Analysis", level=1)
    _make_para(doc, "Method: {} correlation.".format(corr_result.get("method", "pearson").capitalize()), size=10)
    _make_para(doc, corr_result.get("summary", ""), size=9)
    corr_df = corr_result.get("matrix", pd.DataFrame())
    if not corr_df.empty:
        _make_heading(doc, "6.1 Correlation Matrix", level=2)
        display = corr_df.round(3).reset_index()
        display.rename(columns={"index": "Variable"}, inplace=True)
        _add_dataframe_table(doc, display)
    if corr_result.get("heatmap_png"):
        _make_heading(doc, "6.2 Correlation Heatmap", level=2)
        _add_image_from_bytes(doc, corr_result["heatmap_png"], 5.0)
    _make_para(doc, "Variables with |r| > 0.7 indicate strong correlation; 0.3 < |r| < 0.7 indicate moderate; "
               "|r| < 0.3 indicate weak or negligible correlation.", size=9)
    doc.add_page_break()

    # ===== 7. REGRESSION =====
    _make_heading(doc, "7. Regression Analysis", level=1)
    if reg_result and reg_result.get("success"):
        _make_para(doc, "Target Variable: {}".format(target_col))
        _make_para(doc, "R-squared: {}, Adjusted R-squared: {}".format(
            reg_result["r_squared"], reg_result["adj_r_squared"]))
        _make_para(doc, "F-statistic: {}, p-value: {}".format(reg_result["f_stat"], reg_result["f_pvalue"]))
        _make_heading(doc, "7.1 Coefficient Estimates", level=2)
        if not reg_result.get("coefficients", pd.DataFrame()).empty:
            _add_dataframe_table(doc, reg_result["coefficients"])
        _make_heading(doc, "7.2 Interpretation", level=2)
        _make_para(doc, reg_result.get("interpretation", ""), size=9)
        if reg_result.get("residual_plot_png"):
            _add_image_from_bytes(doc, reg_result["residual_plot_png"], 5.5)

    # Multi-model comparison
    if reg_compare_result and reg_compare_result.get("success"):
        _make_heading(doc, "7.3 Multi-Model Comparison", level=2)
        comp_df = pd.DataFrame(reg_compare_result.get("metrics", []))
        if not comp_df.empty:
            _add_dataframe_table(doc, comp_df)
        _make_para(doc, "Best model: {} (by cross-validated R-squared)".format(
            reg_compare_result.get("best_model", "N/A")), bold=True, size=10)
        for key, png in reg_compare_result.get("charts", {}).items():
            _add_image_from_bytes(doc, png, 5.0)
    else:
        _make_para(doc, "Regression could not be performed: {}".format(
            reg_result.get("error", "Unknown error") if reg_result else "No result"), size=9)
    doc.add_page_break()

    # ===== 8. CLASSIFICATION =====
    _make_heading(doc, "8. Classification Analysis", level=1)
    if class_result and class_result.get("success"):
        _make_para(doc, "Models evaluated: {} | Classes: {} | Best: {}".format(
            len(class_result.get("metrics", [])), class_result.get("n_classes", 0),
            class_result.get("best_model", "N/A")), bold=True, size=10)
        _make_heading(doc, "8.1 Model Performance Comparison", level=2)
        comp = pd.DataFrame(class_result.get("metrics", []))
        if not comp.empty:
            _add_dataframe_table(doc, comp)
        for key, b64 in class_result.get("charts", {}).items():
            if b64:
                _add_image_from_bytes(doc, base64.b64decode(b64) if isinstance(b64, str) else b64, 5.0)
    else:
        _make_para(doc, "Classification not applicable: {}".format(
            class_result.get("error", "") if class_result else "Target is continuous or no target selected."), size=9)
    doc.add_page_break()

    # ===== 9. HYPOTHESIS TESTING =====
    _make_heading(doc, "9. Hypothesis Testing", level=1)
    if hyp_result:
        _make_para(doc, hyp_result.get("summary", ""), size=10)
        if hyp_result.get("t_tests"):
            _make_heading(doc, "9.1 Independent t-Tests with Cohen's d", level=2)
            _add_dataframe_table(doc, pd.DataFrame(hyp_result["t_tests"]))
        if hyp_result.get("anovas"):
            _make_heading(doc, "9.2 One-way ANOVA with Eta-squared", level=2)
            _add_dataframe_table(doc, pd.DataFrame(hyp_result["anovas"]))
        if hyp_result.get("chi_squares"):
            _make_heading(doc, "9.3 Chi-Square Tests with Cramer's V", level=2)
            _add_dataframe_table(doc, pd.DataFrame(hyp_result["chi_squares"]))
    doc.add_page_break()

    # ===== 10. FEATURE ENGINEERING =====
    _make_heading(doc, "10. Feature Importance Analysis", level=1)
    if fi_result and fi_result.get("success"):
        methods = fi_result.get("methods", {})
        if methods.get("random_forest"):
            _make_heading(doc, "10.1 Random Forest Importance", level=2)
            _add_dataframe_table(doc, pd.DataFrame(methods["random_forest"]))
        if methods.get("mutual_information"):
            _make_heading(doc, "10.2 Mutual Information", level=2)
            _add_dataframe_table(doc, pd.DataFrame(methods["mutual_information"]))
        if methods.get("target_correlation"):
            _make_heading(doc, "10.3 Target Correlation", level=2)
            _add_dataframe_table(doc, pd.DataFrame(methods["target_correlation"]))
        for key, b64 in fi_result.get("charts", {}).items():
            if b64:
                _add_image_from_bytes(doc, base64.b64decode(b64) if isinstance(b64, str) else b64, 5.0)
    else:
        _make_para(doc, "Feature importance analysis could not be performed.", size=9)
    doc.add_page_break()

    # ===== 11. CLUSTERING =====
    _make_heading(doc, "11. Clustering Analysis", level=1)
    if cluster_result and cluster_result.get("success"):
        _make_para(doc, "Method: {} | Optimal clusters: {} | Silhouette: {:.3f} | "
                   "Davies-Bouldin: {}".format(cluster_result.get("method", "kmeans"),
                                               cluster_result["best_k"],
                                               max(s["silhouette"] for s in cluster_result["silhouette_scores"]),
                                               cluster_result.get("db_score", "N/A")), bold=True, size=10)
        _make_heading(doc, "11.1 Cluster Profiles", level=2)
        profiles = pd.DataFrame(cluster_result.get("profiles", []))
        if not profiles.empty:
            _add_dataframe_table(doc, profiles)
        for key, b64 in cluster_result.get("charts", {}).items():
            if b64:
                _add_image_from_bytes(doc, base64.b64decode(b64) if isinstance(b64, str) else b64, 5.0)
    else:
        _make_para(doc, "Clustering not applicable: {}".format(
            cluster_result.get("error", "") if cluster_result else "Insufficient data."), size=9)
    doc.add_page_break()

    # ===== 12. LIMITATIONS & RISKS =====
    _make_heading(doc, "12. Limitations and Risk Management", level=1)
    limitations = [
        ("Data Quality", "Missing values and outliers may bias results.", "Imputation and outlier handling applied where appropriate."),
        ("Model Scope", "Linear models assume linear relationships.", "Non-linear models (Random Forest, Gradient Boosting) included for comparison."),
        ("Generalizability", "Results reflect the specific dataset analyzed.", "Cross-validation used to estimate out-of-sample performance."),
        ("Causality", "Correlation and regression do not establish causation.", "Results should be treated as associations, not causal claims."),
    ]
    lim_table = doc.add_table(rows=len(limitations) + 1, cols=3, style="Light Grid Accent 1")
    lim_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(["Limitation", "Impact", "Mitigation"]):
        lim_table.rows[0].cells[j].text = h
        for p in lim_table.rows[0].cells[j].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for i, (lim, imp, mit) in enumerate(limitations):
        for j, val in enumerate([lim, imp, mit]):
            lim_table.rows[i + 1].cells[j].text = val
            for p in lim_table.rows[i + 1].cells[j].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)
    doc.add_page_break()

    # ===== 13. RECOMMENDATIONS =====
    _make_heading(doc, "13. Conclusions and Recommendations", level=1)

    rec_parts = []
    if problem_statement:
        rec_parts.append("**Problem Addressed:** {}".format(problem_statement[:250]))
        rec_parts.append("")
    rec_parts.append(
        "This analysis of the \"{}\" dataset ({:,} observations, {} variables) reveals the following key findings:".format(
            dataset_name, rows, cols))

    if reg_result and reg_result.get("success"):
        rec_parts.append(
            "1. The regression model explains {:.1%} of variance in {}. Significant predictors include those with p < 0.05. "
            "Further investigation of interaction effects and non-linear transformations may improve model fit.".format(
                reg_result["r_squared"], target_col))

    if class_result and class_result.get("success") and class_result.get("metrics"):
        rec_parts.append(
            "2. {} achieved the best classification performance (F1 = {:.3f}). "
            "For production deployment, consider the accuracy-speed trade-off between Random Forest and Gradient Boosting.".format(
                class_result["best_model"], class_result["metrics"][0].get("f1_score", 0)))

    if hyp_result:
        sig_count = hyp_result.get("significant_count", 0)
        rec_parts.append(
            "3. {} statistically significant findings emerged from hypothesis testing. "
            "These associations warrant deeper investigation through controlled experiments.".format(sig_count))

    if cluster_result and cluster_result.get("success"):
        rec_parts.append(
            "4. The data naturally segments into {} clusters. Targeted strategies per segment could improve outcomes -- "
            "for example, high-value segments may benefit from differentiated treatment.".format(cluster_result["best_k"]))

    rec_parts.append(
        "5. Recommendation: For future analyses, consider collecting additional data on identified key predictors, "
        "performing longitudinal analysis to track changes over time, and validating findings on an independent holdout sample.")

    for part in rec_parts:
        _make_para(doc, part, size=10)

    # ----- Save -----
    safe_name = "".join(c if c.isalnum() or c in "._- " else "_" for c in dataset_name)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_path = OUTPUT_DIR / "report_{}_{}.docx".format(safe_name, timestamp)
    doc.save(str(out_path))
    return str(out_path)

# ===== Added in v3: Method Development chapter =====

def _add_method_development(doc, auto_ml_result, preprocess_log):
    """Add the Method Development chapter to the report."""
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    _make_heading(doc, "Method Development", level=1)

    if auto_ml_result and auto_ml_result.get("success"):
        _make_para(doc, auto_ml_result.get("narrative", ""), size=10)

        # Model comparison table
        metrics = auto_ml_result.get("metrics", [])
        if metrics:
            _make_heading(doc, "Model Performance Comparison", level=2)
            import pandas as pd
            comp_df = pd.DataFrame(metrics)
            if not comp_df.empty:
                _add_dataframe_table(doc, comp_df)

        # Charts
        for key, b64 in auto_ml_result.get("charts", {}).items():
            if b64:
                import base64
                _add_image_from_bytes(doc, base64.b64decode(b64), 5.0)

        _make_para(doc, f"Best model: {auto_ml_result.get('best_model', 'N/A')} "
                   f"(score: {auto_ml_result.get('best_score', 0):.4f}). "
                   f"Analysis completed in {auto_ml_result.get('elapsed_seconds', 0):.1f} seconds.", bold=True, size=10)

    # Preprocessing log
    if preprocess_log:
        _make_heading(doc, "Data Preprocessing Decisions", level=2)
        for entry in preprocess_log:
            phase = entry.get("phase", "")
            action = entry.get("action", "")
            result = entry.get("result", "")
            chosen = entry.get("chosen", "")
            strategies = entry.get("strategies_tested", [])

            _make_para(doc, f"[{phase}] {action.upper()}", bold=True, size=9)
            if chosen:
                _make_para(doc, f"  Chosen: {chosen}", size=8)
            _make_para(doc, f"  {result}", size=8)

            if strategies:
                for s in strategies:
                    name = s.get("name", "unknown")
                    if "error" in s:
                        _make_para(doc, f"    - {name}: FAILED ({s['error']})", size=8)
                    else:
                        details = ", ".join(f"{k}={v}" for k, v in s.items() if k != "name")
                        _make_para(doc, f"    - {name}: {details}", size=8)
        doc.add_page_break()

    # Limitations
    _make_heading(doc, "Limitations & Future Work", level=1)
    limitations = [
        "All models were evaluated on a single train/test split. Cross-validation provides stability estimates but does not guarantee out-of-sample performance on new data distributions.",
        "The auto-ML pipeline selects the best model based on the chosen metric (F1 for classification, R-squared for regression). Other metrics may suggest different optimal models.",
        "Hyperparameter tuning was limited to default values. Grid search or Bayesian optimization could further improve performance.",
        "Feature engineering was limited to basic encoding. Domain-specific feature transformations could unlock additional predictive power.",
        "The ensemble strategy combined top-3 models via voting. More sophisticated stacking with meta-learners may yield better results.",
    ]
    for i, lim in enumerate(limitations):
        _make_para(doc, f"{i+1}. {lim}", size=9)

    # Future work
    _make_heading(doc, "Future Work", level=2)
    future_items = [
        "Implement automated hyperparameter tuning via Bayesian optimization (Optuna or Hyperopt).",
        "Add feature engineering pipeline: polynomial interactions, target encoding, dimensionality reduction (PCA, t-SNE).",
        "Integrate deep learning models (MLP, TabNet) for large datasets where traditional ML may underfit.",
        "Add model interpretability via SHAP values and LIME explanations.",
        "Implement automated bias detection and fairness metrics for classification tasks.",
    ]
    for item in future_items:
        _make_para(doc, f"- {item}", size=9)
